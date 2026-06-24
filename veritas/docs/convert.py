"""Convert EXECUTIVE_SUMMARY.html to PDF (Playwright) and DOCX (python-docx)."""
import json, os, re, pathlib, textwrap

# ── PDF with Playwright ──────────────────────────────────────────────────────
def to_pdf(html_path: str, pdf_path: str):
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.goto(f"file://{pathlib.Path(html_path).resolve()}")
        page.wait_for_load_state("networkidle")
        page.pdf(path=pdf_path, format="A4", print_background=True,
                 margin={"top": "0.5in", "bottom": "0.5in", "left": "0.6in", "right": "0.6in"})
        browser.close()
    print(f"  -> PDF -> {pdf_path}")

# ── DOCX conversion ──────────────────────────────────────────────────────────
def _style(doc, oxblood="#5c1a1a", gold="#a67c2f", text="#1a1a1a", muted="#6b5e4a"):
    """Apply a consistent style to the document."""
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = 11 * 12700  # 11pt in EMU
    style.paragraph_format.space_after = 6 * 12700
    style.paragraph_format.line_spacing = 1.5
    # Heading styles
    for level, (sz, bold, clr) in {
        0: (22, True, gold),   # Title
        1: (16, True, muted),  # H2
        2: (13, True, text),   # H3
        3: (11, True, gold),   # H4
    }.items():
        s = doc.styles[f"Heading {level + 1}" if level < 4 else "Heading 4"]
        s.font.name = "Georgia"
        s.font.size = sz * 12700
        s.font.bold = bold
        s.font.color.rgb = _rgb(clr)
        s.paragraph_format.space_before = 12 * 12700 * (level + 1)
        s.paragraph_format.space_after = 4 * 12700

def _rgb(hex_str):
    from docx.shared import RGBColor
    h = hex_str.lstrip("#")
    return RGBColor(*[int(h[i:i+2], 16) for i in (0, 2, 4)])

def _add_metrics(doc, rows):
    """Add a 4-column metrics row as a table."""
    from docx.shared import Inches, Pt
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, (num, label) in enumerate(rows):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(num))
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = _rgb("#a67c2f")
        run.font.name = "Georgia"
        p2 = cell.add_paragraph(label)
        p2.runs[0].font.size = Pt(7)
        p2.runs[0].font.color.rgb = _rgb("#6b5e4a")
        p2.runs[0].font.name = "Georgia"

def _add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = _rgb("#8b6d2a")
    # Body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = _rgb("#6b5e4a")
    # styling
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = 2 * 12700
                p.paragraph_format.space_after = 2 * 12700

def to_docx(html_path: str, docx_path: str):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _style(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = st.add_run("TRUTHSEEKERS")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = _rgb("#a67c2f")
    
    st2 = doc.add_paragraph()
    st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = st2.add_run("Engine for Evidence-Grounded Knowledge")
    run2.font.size = Pt(11)
    run2.font.color.rgb = _rgb("#6b5e4a")
    run2.italic = True

    doc.add_paragraph()

    # ── Lede ──
    doc.add_paragraph(
        "Truthseekers is a knowledge construction engine — it produces structured, evidence-grounded "
        "encyclopedia articles across all truth categories (confirmed, contested, suppressed, speculative), "
        "exposes framing and structural distortion, and refuses to treat the absence of evidence as silence."
    ).runs[0].italic = True

    _add_metrics(doc, [
        ("3", "Strict Epistemic Layers"), ("9", "Pipeline Nodes"),
        ("14", "Agent Tools"), ("23", "Storage CRUD Ops"),
    ])

    # ── Sections ──
    sections = [
        ("The Problem", [
            "Every knowledge platform today is a chatbot with a search query — fluent but flat, unable to "
            "distinguish consensus from suppression. Conventional systems cannot separate evidence quality "
            "from narrative fluency, detect missing evidence as signal, trace provenance, or surface "
            "suppressed evidence. Truthseekers solves all six by design: strict three-layer epistemic "
            "separation, gap detection as a first-class signal, dual-output language mapping, six-axis "
            "confidence vectors (evidence strength, corroboration, source diversity, recency, contradiction "
            "level, bias risk), and four explicit truth categories."
        ]),
        ("Architecture in Brief", [
            "The system has three layers: a Go HTTP orchestrator (stdlib net/http, no frameworks, port 4097) "
            "with a native agent loop (25 iterations, 14 tools, 90k token budget) and a DAG engine; Python "
            "workers (9 stateless epistemic nodes over a JSON subprocess bridge with 4-provider LLM fallback "
            "and mock mode); and a Next.js 15 frontend with React 19, Framer Motion, TanStack Query, and "
            "an editorial design system. Full zero-infrastructure local development.",
            "Zero-framework Go backend on port 4097. 9 Python nodes with shared LLM client and 4-provider "
            "fallback chain (Groq → DO Inference → NVIDIA NIM → OpenAI). React 19 frontend with editorial "
            "gold-and-ink design, structured block renderer, SSE streaming, 3D maps, interactive timelines. "
            "Full local demo from go run ./cmd/server — no Docker, no cloud account, no API keys required."
        ]),
        ("Technical Differentiation", [
            "Strict epistemic layers: three layers with typed I/O contracts, enumerated permissions, and "
            "hard prohibitions at code level. - Claim Graph architecture: articles are generated views over "
            "an atomic claim graph with six-axis confidence vectors. - Gap detection: identifies structural "
            "evidence gaps as a meaningful epistemic signal. - Language mapping: dual-channel output preserving "
            "original framing alongside precision upgrades. - Multi-provider LLM routing with graceful "
            "degradation. - Zero-infrastructure mock mode: every dependency has a fallback."
        ]),
        ("Market Position", []),  # table is added separately
        ("Current Status & Trajectory", [
            "Go backend: live. Python workers: complete. Frontend: live, with SSE streaming fixed, "
            "loading/error states implemented, block rendering fixed, and React.memo optimizations. "
            "DAG-to-HTTP integration is the remaining engineering milestone (standard task with known interfaces). "
            "Infrastructure, testing, and documentation are well-established.",
            "The system works end-to-end today. The DAG pipeline produces valid article output. "
            "The agent holds conversations with 14 tools. The frontend streams real-time progress. "
            "The remaining integration is a standard engineering task with known interfaces on both sides."
        ]),
        ("Technology Stack", [
            "Go 1.22+ with stdlib net/http, MongoDB with in-memory mock, HMAC-SHA256 JWT auth, "
            "OpenAI-compatible API (Groq + DO Inference), Python 3.11+ subprocesses. "
            "Frontend: Next.js 15.5.19, React 19, Tailwind CSS 4, TanStack Query, Framer Motion 11, "
            "Mermaid.js, Leaflet, React Three Fiber.",
            "The Go backend uses zero web frameworks — deliberate choice for small, stable API surface, "
            "reduced attack surface, and sub-10-second compile times."
        ]),
        ("Team", [
            "Solo-built across the full stack: Go systems programming, Python ML/LLM pipelines, and React "
            "frontend architecture. Custom editorial design system built from scratch. Novel three-layer "
            "epistemic architecture with typed I/O contracts and hard prohibitions at code level."
        ]),
        ("Roadmap", [
            "Phase 1 (Architecture Design): Done. Phase 2 (Core Implementation): Done. "
            "Phase 3 (Frontend Integration): ~70% done, 2 weeks remaining. "
            "Phase 4 (Epistemic Tuning): Pending, 3 weeks. Phase 5 (PostgreSQL Migration): Pending, 2 weeks. "
            "Phase 6 (Production Hardening): Pending, 3 weeks."
        ]),
        ("Key Metrics", [
            "~15k lines Go, ~8k lines Python, ~25k lines TypeScript/JS. 19 Go source files, "
            "14 agent tools, 10 Python worker scripts, 23 database CRUD methods, "
            "4 LLM provider gateways, 4 truth categories, 6 confidence axes, "
            "9 DAG node operations, 3 epistemic layers."
        ]),
    ]

    for title, paras in sections:
        doc.add_heading(title, level=1)
        if title == "Market Position":
            _add_table(doc,
                ["Feature", "Truthseekers", "Wikipedia", "ChatGPT/Claude", "Perplexity"],
                [
                    ["Epistemic layer separation", "✓ Code-level", "✗ Editorial only", "✗ None", "✗ None"],
                    ["Evidence gap detection", "✓ First-class", "✗ Not modeled", "✗ Not modeled", "✗ Not modeled"],
                    ["Lang. framing detection", "✓ Dual-channel", "✗ None", "✗ Reproduces", "✗ Reproduces"],
                    ["Multi-dim confidence", "✓ 6-axis", "✗ Binary", "✗ Scalar", "✗ None"],
                    ["Suppressed evidence", "✓ Explicit cat.", "✗ Gated", "✗ Bias", "✗ Bias"],
                    ["Full provenance", "✓ Full chain", "✓ Per sentence", "✗ Hallucinated", "✓ URLs"],
                    ["Open source", "✓ MIT", "✓ CC-BY-SA", "✗ Proprietary", "✗ Proprietary"],
                    ["Local dev", "✓ Zero infra", "✓ Needs stack", "✗ API only", "✗ API only"],
                ],
            )
            doc.add_paragraph(
                "Reference markets: knowledge management ($57B by 2027), AI content generation ($18B by 2028), "
                "research intelligence platforms ($12B by 2026). Truthseekers sits at the intersection."
            )
        elif title == "Key Metrics":
            _add_metrics(doc, [
                ("~15k", "Lines of Go"), ("~8k", "Lines of Python"),
                ("~25k", "Lines of TS/JS"), ("19", "Go Source Files"),
            ])
            _add_metrics(doc, [
                ("14", "Agent Tools"), ("10", "Python Scripts"),
                ("23", "DB CRUD Methods"), ("4", "LLM Providers"),
            ])
            _add_metrics(doc, [
                ("4", "Truth Categories"), ("6", "Confidence Axes"),
                ("9", "DAG Nodes"), ("3", "Epistemic Layers"),
            ])
        else:
            for p in paras:
                doc.add_paragraph(p)

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Truthseekers - System Overview - June 2026")
    fr.font.size = Pt(9)
    fr.font.color.rgb = _rgb("#6b5e4a")

    doc.save(docx_path)
    print(f"  -> DOCX -> {docx_path}")


if __name__ == "__main__":
    root = pathlib.Path(__file__).parent
    html_path = root / "EXECUTIVE_SUMMARY.html"
    pdf_path = root / "EXECUTIVE_SUMMARY.pdf"
    docx_path = root / "EXECUTIVE_SUMMARY.docx"

    to_pdf(str(html_path), str(pdf_path))
    to_docx(str(html_path), str(docx_path))
    print("Done.")
